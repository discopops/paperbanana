"""DOCX (Microsoft Word) document loader with section extraction."""

from __future__ import annotations

import re
from pathlib import Path
from typing import ClassVar, Optional

from paperbanana.core.types import LoadResult, LoaderConfig
from paperbanana.loaders.base import BaseLoader


class DOCXLoader(BaseLoader):
    """Loader for Microsoft Word documents (.docx).

    Extracts text from paragraphs and uses heading styles to identify sections.

    Requires: python-docx>=1.0 (optional dependency)
    """

    format_name: ClassVar[str] = "Microsoft Word"
    supported_extensions: ClassVar[list[str]] = [".docx"]

    # Methodology section patterns (case-insensitive)
    METHODOLOGY_PATTERNS = [
        r"method(?:ology|s)?",
        r"approach(?:es)?",
        r"experimental\s+(?:setup|design|methodology)",
        r"materials?\s+and\s+methods?",
        r"proposed\s+(?:method|approach)",
        r"framework",
    ]

    async def load(self, path: Path, config: LoaderConfig | None = None) -> LoadResult:
        """Load content from a DOCX file.

        Args:
            path: Path to the DOCX file
            config: Optional configuration for extraction behavior

        Returns:
            LoadResult with extracted text and metadata

        Raises:
            FileNotFoundError: If the file doesn't exist
            ImportError: If python-docx is not installed
            ValueError: If the document cannot be parsed
        """
        if config is None:
            config = LoaderConfig()

        if not path.exists():
            raise FileNotFoundError(f"File not found: {path}")

        # Import python-docx
        try:
            from docx import Document
        except ImportError as e:
            raise ImportError(
                "DOCX loading requires python-docx. Install with: pip install paperbanana[docx]"
            ) from e

        # Open document
        try:
            doc = Document(path)
        except Exception as e:
            raise ValueError(f"Failed to open DOCX: {e}") from e

        # Extract full text with paragraph structure
        paragraphs_data = []
        for para in doc.paragraphs:
            # Get paragraph text and style
            text = para.text.strip()
            if not text:
                continue

            # Detect if this is a heading
            style_name = para.style.name if para.style else ""
            is_heading = "Heading" in style_name

            paragraphs_data.append(
                {"text": text, "is_heading": is_heading, "style": style_name}
            )

        # Extract full text
        full_text = "\n\n".join(p["text"] for p in paragraphs_data)

        # Extract methodology section if configured
        source_context = full_text
        extraction_method = "full_text"

        if config.extract_methodology:
            extracted = self._extract_methodology_section(paragraphs_data)
            if extracted:
                source_context = extracted
                extraction_method = "section_extraction"

        # Apply max length if specified
        if config.max_length and len(source_context) > config.max_length:
            source_context = source_context[: config.max_length]

        # Build metadata
        metadata = {
            "format": "docx",
            "file_size": path.stat().st_size,
            "paragraph_count": len(doc.paragraphs),
            "extraction_method": extraction_method,
            "truncated": config.max_length is not None
            and len(source_context) == config.max_length,
        }

        return LoadResult(
            source_context=source_context,
            metadata=metadata,
            images=[],
        )

    def _extract_methodology_section(
        self, paragraphs_data: list[dict]
    ) -> Optional[str]:
        """Extract methodology section using heading detection.

        Args:
            paragraphs_data: List of paragraph dicts with text, is_heading, style

        Returns:
            Extracted methodology section, or None if not found
        """
        # Find methodology heading
        start_idx = None
        heading_style = None

        for i, para in enumerate(paragraphs_data):
            if not para["is_heading"]:
                continue

            text_lower = para["text"].lower()
            for pattern in self.METHODOLOGY_PATTERNS:
                if re.search(pattern, text_lower, re.IGNORECASE):
                    start_idx = i + 1
                    heading_style = para["style"]
                    break

            if start_idx is not None:
                break

        if start_idx is None:
            return None

        # Find next heading of same or higher level (end of section)
        end_idx = len(paragraphs_data)
        for i in range(start_idx, len(paragraphs_data)):
            para = paragraphs_data[i]
            if para["is_heading"] and para["style"] == heading_style:
                end_idx = i
                break

        # Extract paragraphs
        methodology_paras = paragraphs_data[start_idx:end_idx]
        methodology_text = "\n\n".join(p["text"] for p in methodology_paras)

        # Return None if extraction is too short
        if len(methodology_text) < 100:
            return None

        return methodology_text

    def is_available(self) -> bool:
        """Check if python-docx is installed."""
        try:
            import docx  # noqa: F401

            return True
        except ImportError:
            return False

    def get_install_hint(self) -> str:
        """Get installation instructions for python-docx."""
        return "Install with: pip install paperbanana[docx]"
